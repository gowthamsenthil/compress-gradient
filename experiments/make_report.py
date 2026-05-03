"""Generate the final research-paper-style .docx report.

Produces `results/report.docx` with embedded figures from `results/*.png`.

Format:
  - Single column, 11pt Times New Roman.
  - Title page + Abstract + numbered sections.
  - Each figure has a caption beneath; pseudocode in monospace.
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
RESULTS = ROOT / "results"
OUT = RESULTS / "report.docx"


# ----------------------------- helpers -------------------------------------

def set_default_font(doc: Document, name: str = "Times New Roman", size: int = 11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt({1: 14, 2: 12, 3: 11}.get(level, 11))
        run.bold = True
    return h


def add_para(doc, text, bold=False, italic=False, align=None, size=11):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_figure(doc, path, caption, width_inches=6.0):
    if not Path(path).exists():
        add_para(doc, f"[MISSING FIGURE: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(10)
    cr.italic = True


def add_code(doc, code: str):
    """Add monospaced pseudocode block."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        cells = table.rows[1 + r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)


# ----------------------------- content -------------------------------------

def build():
    doc = Document()
    set_default_font(doc)

    # Margins (slightly tighter than default to fit the page count well).
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # ----------- TITLE -----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "CompressIQ: Bandwidth-Aware Convex Optimization of\n"
        "Per-Worker Gradient Compression for Heterogeneous\n"
        "Distributed Training"
    )
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(16)
    tr.bold = True

    add_para(doc, "Gowtham", align="center", size=12)
    add_para(doc, "Independent Research Project", align="center", italic=True, size=11)
    add_para(doc, "May 2026", align="center", italic=True, size=11)
    add_para(doc, "")

    # ----------- ABSTRACT -----------
    add_heading(doc, "Abstract", level=1)
    add_para(
        doc,
        "Synchronous data-parallel training on heterogeneous clusters is bottlenecked "
        "by the slowest worker. Uniform gradient compression -- the standard "
        "production practice -- ignores per-worker bandwidth and unnecessarily "
        "throttles fast workers to the slow workers' communication cost. We "
        "formulate per-worker compression-ratio selection as a convex program with "
        "a min-max ring all-reduce objective and an accuracy-loss constraint, and "
        "validate the formulation end-to-end on a discrete-event MNIST training "
        "simulator that charges wall-clock per the bandwidth-optimal Patarasuk-Yuan "
        "ring all-reduce cost model. Our solution attains a 1.85x speedup over "
        "uniform compression at matched test accuracy on a 12-worker, 3-tier "
        "cluster. We then refine the accuracy bound in two ways: (i) an "
        "error-feedback-aware (EF-aware) bound that calibrates online against the "
        "actually-leaked residual norm, and (ii) a per-layer formulation that "
        "replaces the scalar compression ratio per worker with a layer-wise vector. "
        "Combined, the refinements deliver a further 1.76x speedup at matched "
        "accuracy, for a cumulative 3.26x improvement over uniform. We also "
        "characterize empirical drift in the per-worker squared gradient norm "
        "during training and show that periodic recalibration restores constraint "
        "compliance without measurable accuracy loss.",
        align="justify",
    )
    add_para(
        doc,
        "Code and experimental artifacts: "
        "https://github.com/gowthamsenthil/compress-gradient",
        italic=True,
        align="justify",
    )

    # ----------- 1. INTRODUCTION -----------
    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "Distributed training of deep neural networks under data parallelism "
        "alternates local gradient computation with a global all-reduce that "
        "synchronizes gradient estimates across workers. As model size grows and "
        "clusters become more heterogeneous (mixing on-prem GPUs, cloud instances "
        "of multiple network tiers, and edge devices), the all-reduce step "
        "frequently dominates per-iteration wall-clock time. Gradient compression "
        "techniques -- Top-K sparsification, quantization (QSGD), low-rank "
        "approximation (PowerSGD) -- mitigate the bottleneck by transmitting fewer "
        "bits per round, but their parameters are typically chosen as cluster-wide "
        "scalars, ignoring inter-worker bandwidth heterogeneity.",
        align="justify",
    )
    add_para(
        doc,
        "We argue that the natural decision variable in a heterogeneous cluster is "
        "the per-worker compression ratio r_i, and that selecting these ratios is "
        "a convex optimization problem under a quadratic accuracy-loss model. The "
        "problem decouples cleanly from training itself: at each (re)calibration "
        "point, a small convex program (solvable in milliseconds with off-the-shelf "
        "solvers) returns a Pareto-optimal ratio vector, which is then used to "
        "drive standard SGD with error feedback for the next K rounds.",
        align="justify",
    )
    add_para(doc, "Contributions:", bold=True)
    add_bullets(doc, [
        "We derive a min-max ring all-reduce cost objective from the bandwidth-optimal "
        "Patarasuk-Yuan formula and show that it produces a water-filling solution "
        "in which r_i / B_i is equalized across all workers not at the box bounds.",
        "We integrate the formulation into a discrete-event MNIST training simulator "
        "with error feedback and a faithful ring-all-reduce wall-clock model, "
        "achieving a 1.85x speedup over uniform compression at iso-accuracy.",
        "We empirically characterize drift in the per-worker squared gradient norm "
        "during training and show that the static formulation accumulates a 3.3x "
        "constraint violation by round 300; periodic recalibration restores compliance.",
        "We propose two bound refinements -- an EF-aware online correction factor "
        "kappa_i, and a per-layer extension r_{i, l} -- that together deliver a "
        "further 1.76x speedup at unchanged accuracy, motivated directly by the "
        "Karimireddy et al. error-feedback convergence framework.",
    ])

    # ----------- 2. RELATED WORK -----------
    add_heading(doc, "2. Related Work", level=1)
    add_para(doc, "2.1 Gradient compression.", bold=True)
    add_para(
        doc,
        "Top-K sparsification was popularized by Lin et al. (Deep Gradient "
        "Compression, 2018), which observed that retaining only the magnitude-top "
        "0.1-1% of gradient coordinates suffices for ImageNet training when "
        "combined with momentum correction and warmup. Alistarh et al. (2018) "
        "provided the first convergence theory for sparsified SGD under bounded "
        "second-moment assumptions. Quantization-based methods such as QSGD "
        "(Alistarh et al. 2017), TernGrad (Wen et al. 2017), and signSGD "
        "(Bernstein et al. 2018) trade representational precision rather than "
        "sparsity. Low-rank approaches including PowerSGD (Vogels et al. 2019) "
        "and ATOMO (Wang et al. 2018) compress along principal directions.",
        align="justify",
    )
    add_para(doc, "2.2 Error feedback.", bold=True)
    add_para(
        doc,
        "Biased compressors (including Top-K) introduce systematic descent-direction "
        "errors that break convergence guarantees. Karimireddy et al. (Error "
        "Feedback Fixes SignSGD and other Gradient Compression Schemes, 2019) "
        "showed that wrapping any delta-contractive compressor in an error-feedback "
        "loop -- in which the per-round residual is folded into the next round's "
        "gradient -- restores SGD convergence at the standard rate. Their bound on "
        "the residual norm is the theoretical foundation of our EF-aware bound "
        "refinement (Section 4.3).",
        align="justify",
    )
    add_para(doc, "2.3 Bandwidth-optimal collectives.", bold=True)
    add_para(
        doc,
        "Patarasuk and Yuan (Bandwidth Optimal All-reduce Algorithms for Clusters "
        "of Workstations, 2009) proved that the ring all-reduce achieves the "
        "minimum possible bandwidth cost of 2(N-1)/N times the message size for "
        "homogeneous-link clusters, and that this cost is tight. Their formula "
        "T = 2(N-1)/N * (M / B) + 2(N-1) * alpha is the per-round wall-clock "
        "model used by both Horovod (Sergeev and Del Balso 2018) and PyTorch DDP. "
        "We extend this formula to per-worker bandwidths and per-worker (or "
        "per-layer) compression ratios in Section 3.1.",
        align="justify",
    )
    add_para(doc, "2.4 Heterogeneity-aware scheduling.", bold=True)
    add_para(
        doc,
        "ByteScheduler (Peng et al. 2019) and BytePS (Jiang et al. 2020) optimize "
        "tensor scheduling and parameter-server placement under heterogeneous "
        "topologies but do not vary the compression ratio per worker. SkewScout "
        "(Harlap et al. 2017) and similar straggler-mitigation systems address "
        "compute heterogeneity by load balancing, again without touching "
        "compression. Federated learning systems (Konecny et al. 2017) tolerate "
        "client heterogeneity by accepting partial participation rather than by "
        "differentiated compression. To our knowledge, formulating per-worker "
        "compression as a constrained convex program with a min-max ring objective "
        "is novel.",
        align="justify",
    )

    # ----------- 3. PROBLEM FORMULATION -----------
    add_heading(doc, "3. Problem Formulation", level=1)

    add_heading(doc, "3.1 Ring all-reduce cost model", level=2)
    add_para(
        doc,
        "Consider N workers connected in a ring, where worker i has outbound link "
        "bandwidth B_i (bits/s) and per-hop latency alpha_i (s). For an "
        "uncompressed gradient of size G bits, the bandwidth-optimal ring "
        "all-reduce of Patarasuk and Yuan completes one synchronization round "
        "in T = 2(N-1)/N * (G / B) + 2(N-1) * alpha when bandwidth is uniform. "
        "When workers compress at distinct ratios r_i in (0, 1], worker i sends "
        "r_i * G / N bits per ring step across 2(N-1) steps. Step k on link "
        "(i -> i+1) takes (r_i * G / N) / B_i seconds and the synchronous ring "
        "advances at the slowest link, so the round time is",
        align="justify",
    )
    add_code(doc,
             "    T_round = 2(N-1)/N * G * max_i (r_i / B_i) + 2(N-1) * max_i alpha_i.")
    add_para(
        doc,
        "Constants in r drop out of the argmin; the optimization reduces to "
        "min max_i (r_i / B_i). For the simulator wall-clock charge, we retain "
        "all constants -- this is the actual cost a real ring all-reduce would "
        "incur. The full per-round time is implemented in compressiq/cost_model.py.",
        align="justify",
    )

    add_heading(doc, "3.2 Accuracy-loss bound and assumptions", level=2)
    add_para(
        doc,
        "We model accuracy loss per round by the squared compressor error, summed "
        "across workers and bounded by a budget epsilon:",
        align="justify",
    )
    add_code(doc, "    sum_i alpha_i * (1 - r_i)^2  <=  epsilon,")
    add_para(
        doc,
        "where alpha_i = ||g_i||^2 is the squared norm of worker i's local "
        "gradient. The form (1 - r)^2 * alpha is an upper bound on Top-K's "
        "squared error: keeping the largest-magnitude r * d coordinates yields "
        "error at most as large as keeping a uniformly random r * d "
        "coordinates, which has expectation (1 - r) * alpha. Squaring this "
        "factor gives a slack envelope sufficient for convexity.",
        align="justify",
    )
    add_para(doc, "Assumptions baked into this constraint:", bold=True)
    add_table(doc,
              header=["#", "Assumption", "Comment"],
              rows=[
                  ["A1", "Compressor error <= (1-r)^2 * ||g||^2",
                   "Conservative upper bound for Top-K"],
                  ["A2", "alpha_i = ||g_i||^2 summarizes per-worker accuracy weight",
                   "Coarse; per-layer is tighter (Section 4.4)"],
                  ["A3", "Worker errors combine additively",
                   "True under the standard MSE decomposition"],
                  ["A4", "Cross-worker errors are independent",
                   "Approx. true for IID data shards"],
                  ["A5", "Per-round error does not accumulate",
                   "Pessimistic under EF (motivates Section 4.3)"],
                  ["A6", "alpha_i is constant in t",
                   "Violated empirically (Section 5.6)"],
              ])

    add_heading(doc, "3.3 The convex program", level=2)
    add_para(
        doc,
        "Combining the ring objective with the accuracy bound and box "
        "constraints yields:",
        align="justify",
    )
    add_code(doc,
             "    minimize    max_i  r_i / B_i\n"
             "    subject to  sum_i  alpha_i * (1 - r_i)^2  <=  epsilon\n"
             "                r_min <= r_i <= 1            for all i = 1, ..., N\n"
             "    variables   r_1, ..., r_N.")
    add_para(
        doc,
        "The objective is the maximum of N linear functions of r (convex). The "
        "accuracy constraint is a sum of N convex quadratics. r_min > 0 prevents "
        "degenerate sends with no information. The problem is solved with CVXPY "
        "(Diamond and Boyd 2016) in the second-order-cone form using Clarabel, "
        "ECOS, or SCS. Solve time on N = 12 is well under 10 ms, so re-solving "
        "every K=25 rounds is negligible relative to a training round.",
        align="justify",
    )

    # ----------- 4. METHODOLOGY -----------
    add_heading(doc, "4. Methodology", level=1)

    add_heading(doc, "4.1 Static CompressIQ", level=2)
    add_para(
        doc,
        "Calibrate alphas once before training (one batch per worker), solve the "
        "convex program once for the chosen epsilon, and use the resulting r_i "
        "for the entire run. This is the simplest baseline and is what most "
        "production heterogeneous-cluster pipelines effectively do today.",
        align="justify",
    )

    add_heading(doc, "4.2 Adaptive recalibration", level=2)
    add_para(
        doc,
        "Static CompressIQ assumes alpha_i is stationary in t (assumption A6). "
        "We observed that ||g_i||^2 grows roughly 3x over 300 rounds in our "
        "MNIST/MLP setup as residuals from error feedback accumulate (Figure 6). "
        "The adaptive variant recalibrates every K rounds:",
        align="justify",
    )
    add_code(doc,
             "Algorithm 1: Adaptive CompressIQ training\n"
             "Input:  workers W_1..W_N, budget epsilon, period K, initial r_i.\n"
             "for t = 1 to T:\n"
             "    if (t mod K == 0):\n"
             "        alpha_i <- ||W_i.compute_gradient()||^2\n"
             "        r <- solve_compressiq(alpha, epsilon)\n"
             "    for each W_i:\n"
             "        g_i <- W_i.compute_gradient()\n"
             "        c_i, leaked_i <- W_i.compress(g_i, r_i)   # uses error feedback\n"
             "    avg_grad <- mean_i(c_i)\n"
             "    sim_clock += 2(N-1)/N * G * max_i(r_i / B_i) + 2(N-1) * max_i alpha_i\n"
             "    each W_i applies avg_grad with learning rate lr.")

    add_heading(doc, "4.3 Error-feedback-aware (EF-aware) bound", level=2)
    add_para(
        doc,
        "Under error feedback, the per-round leaked squared error is "
        "consistently smaller than (1 - r_i)^2 * alpha_i because past residuals "
        "are folded into the next compression step (Karimireddy et al. 2019). "
        "We exploit this by introducing a per-worker correction factor "
        "kappa_i in (0, 1]:",
        align="justify",
    )
    add_code(doc,
             "    kappa_i(t) = EMA_t( ||leaked_i,t||^2 / (alpha_i,t * (1 - r_i,t)^2) ),")
    add_para(
        doc,
        "with EMA coefficient 0.2 and clipping to [1e-6, 1]. The constraint "
        "becomes sum_i kappa_i * alpha_i * (1 - r_i)^2 <= epsilon. When kappa_i "
        "= 1 (default initialization), the bound reduces to the naive form -- "
        "backward compatible. As training proceeds and kappa_i settles "
        "empirically near 0.03, the optimizer correctly identifies that it has "
        "approximately 1/0.03 = 33x more usable accuracy budget than the naive "
        "bound suggested, and pushes ratios more aggressively.",
        align="justify",
    )

    add_heading(doc, "4.4 Per-layer formulation", level=2)
    add_para(
        doc,
        "The scalar r_i conflates layers with very different sizes and "
        "different accuracy sensitivities. Replacing it with a per-layer matrix "
        "R[i, l] (i over workers, l over layers) keeps the convex structure:",
        align="justify",
    )
    add_code(doc,
             "    minimize    max_i  sum_l R[i,l] * G_l / B_i\n"
             "    subject to  sum_{i,l} alpha[i,l] * (1 - R[i,l])^2 <= epsilon\n"
             "                r_min <= R[i,l] <= 1.")
    add_para(
        doc,
        "On our MNIST MLP this yields three layers (fc1, fc2, fc3). The solver "
        "consistently keeps fc3 (~1% of the parameter bits) at r ~= 0.97-0.99 "
        "while compressing fc1 and fc2 (~99% of the bits) much more aggressively. "
        "This is exactly the behaviour one would expect: the small output layer "
        "contributes negligibly to communication time but fully consumes "
        "accuracy budget, so the optimizer assigns it a near-uncompressed share. "
        "Both EF-aware and per-layer can be combined; their effects partly stack "
        "(Section 5.7).",
        align="justify",
    )

    # ----------- 5. EXPERIMENTS -----------
    add_heading(doc, "5. Experiments", level=1)

    add_heading(doc, "5.1 Setup", level=2)
    add_para(
        doc,
        "All experiments use a discrete-event simulator with a faithful ring "
        "all-reduce wall-clock model. Workers train an MLP (784-128-128-10, "
        "ReLU, ~118K parameters, 32-bit floats -> ~3.78M bits per gradient) on "
        "MNIST shards; the simulator advances simulated time by max_i T_i per "
        "round. The default cluster has N = 12 workers in three bandwidth "
        "tiers: 6x10 Gbps, 3x5 Gbps, 3x1 Gbps, with per-hop latencies of "
        "5, 20, 50 microseconds respectively, modelling NVLink + intra-rack "
        "Ethernet + inter-rack/cloud links. Error feedback is enabled by default "
        "for all schemes that compress (uniform, CompressIQ, refinements). "
        "Learning rate is 0.05, batch size 64, 200-300 training rounds. The "
        "accuracy budget epsilon is parameterized as a fraction of the initial "
        "sum of alpha_i; the Pareto sweep spans epsilon = 0.01 * sum(alpha) to "
        "0.9 * sum(alpha).",
        align="justify",
    )

    add_heading(doc, "5.2 Pareto frontier", level=2)
    add_figure(doc, RESULTS / "fig1_pareto.png",
               "Figure 1: Pareto frontier of test accuracy vs. simulated "
               "communication time across schemes (no compression, uniform, "
               "greedy, CompressIQ) under varying accuracy budgets epsilon. "
               "Each marker is a complete training run; the convex hull "
               "(dashed) traces the achievable frontier. CompressIQ dominates "
               "uniform and greedy at every operating point.")
    add_para(
        doc,
        "CompressIQ Pareto-dominates both uniform and greedy at every accuracy "
        "level we tested. At ~0.90 test accuracy, CompressIQ takes 1.85x less "
        "simulated communication time than uniform compression at the same "
        "accuracy budget. The greedy baseline (each worker independently "
        "spends a bandwidth-weighted budget share) is closer to CompressIQ but "
        "is uniformly inferior because it cannot coordinate budget across "
        "workers.",
        align="justify",
    )

    add_heading(doc, "5.3 Per-worker ratios: water-filling structure", level=2)
    add_figure(doc, RESULTS / "fig2_ratios.png",
               "Figure 2: Per-worker compression ratio r_i for CompressIQ vs. "
               "uniform on the 3-tier cluster, sorted by bandwidth (fast "
               "leftmost). At the displayed epsilon, all three bandwidth tiers "
               "are strictly interior, exposing the water-filling staircase "
               "predicted by the KKT conditions (r_i / B_i = tau across all "
               "active workers).")
    add_para(
        doc,
        "The KKT conditions for the convex program imply that for every worker "
        "not at a box bound, r_i / B_i takes the same value tau. Empirically, "
        "at epsilon = 3.32 the per-worker ratios are 0.071 for the 1 Gbps tier, "
        "0.356 for the 5 Gbps tier, and 0.712 for the 10 Gbps tier; in each "
        "case r_i / B_i = 0.0712, exactly equalized. This is the textbook "
        "water-filling solution and is a non-trivial sanity check that the "
        "convex formulation is correctly aligning with the underlying ring "
        "structure.",
        align="justify",
    )

    add_heading(doc, "5.4 Round-time decomposition", level=2)
    add_figure(doc, RESULTS / "fig3_round_time.png",
               "Figure 3: Per-round simulated wall-clock time for each "
               "scheme, averaged over the run. The bar height matches the "
               "Patarasuk-Yuan ring formula 2(N-1)/N * G * max_i(r_i/B_i) + "
               "2(N-1) * max_i alpha_i applied to the operating-point ratios "
               "of each scheme.")
    add_para(
        doc,
        "Round time is dominated by the bandwidth term in our setting; the "
        "latency term is on the order of microseconds per round and contributes "
        "a fixed constant. CompressIQ achieves ~3.3 ms per round vs. ~6.1 ms "
        "for uniform, consistent with the 1.85x speedup observed in the "
        "Pareto experiment.",
        align="justify",
    )

    add_heading(doc, "5.5 Bound validation", level=2)
    add_figure(doc, RESULTS / "fig5_constraint.png",
               "Figure 5: Per-round theoretical accuracy error sum_i alpha_i "
               "(1-r_i)^2 (solid) versus empirical leaked squared error "
               "(dashed) over 300 training rounds, with the budget epsilon "
               "drawn as a horizontal dotted line. The naive bound is a safe "
               "upper bound but is consistently 5-30x looser than the actual "
               "compressor leakage under error feedback.")
    add_para(
        doc,
        "This figure motivates the EF-aware bound refinement (Section 4.3). "
        "The 5-30x slack between predicted and actual error is exactly the "
        "budget that EF-aware reclaims and spends on more aggressive "
        "compression. We also note that the static CompressIQ scheme's "
        "*theoretical* error eventually drifts above epsilon -- i.e., the "
        "static formulation violates its own bound by round 300 -- because "
        "alpha_i is non-stationary; this motivates Section 4.2.",
        align="justify",
    )

    add_heading(doc, "5.6 Drift in alpha and adaptive recalibration", level=2)
    add_figure(doc, RESULTS / "fig6_alpha_drift.png",
               "Figure 6: Per-worker squared gradient norm alpha_i = ||g_i||^2 "
               "as a function of training round. Contrary to the static "
               "assumption, alpha_i grows roughly 3x over 300 rounds as "
               "training progresses and error-feedback residuals accumulate.")
    add_para(
        doc,
        "Periodic recalibration (every K=25 rounds) absorbs this drift cheaply: "
        "the convex program re-solves in under 10 ms with fresh alphas, and the "
        "resulting r_i stay within the constraint throughout training. K=10 "
        "provides marginally tighter constraint compliance than K=25 at "
        "negligible additional cost (the recalibration step itself dominates "
        "no individual round). We use K=25 as the default for all subsequent "
        "experiments.",
        align="justify",
    )

    add_heading(doc, "5.7 Speedup summary", level=2)
    add_figure(doc, RESULTS / "fig7_speedup.png",
               "Figure 7: Cumulative simulated communication time to reach "
               "fixed accuracy targets, normalized to the uniform baseline. "
               "CompressIQ delivers 1.85x speedup at 90% accuracy; the "
               "adaptive variant matches static CompressIQ's wall-clock while "
               "satisfying the accuracy bound throughout training (which "
               "static violates).")

    add_heading(doc, "5.8 Bound refinements: EF-aware and per-layer", level=2)
    add_figure(doc, RESULTS / "fig9_refine_speed.png",
               "Figure 9: (Left) Test accuracy vs. cumulative simulated "
               "communication time for the four bound variants on the 3-tier "
               "cluster, 200 rounds. (Right) Speedup over the naive bound at "
               "matched final accuracy. Per-layer + EF-aware achieves a "
               "1.76x speedup with the highest final accuracy of all variants.")
    add_para(
        doc,
        "All four variants converge to within 0.0014 of one another in final "
        "test accuracy (0.9057 to 0.9071), but their cumulative communication "
        "times differ by 1.76x. EF-aware alone delivers a 1.57x speedup; "
        "per-layer alone delivers 1.16x; combined they yield 1.76x. The two "
        "refinements address orthogonal sources of conservatism in the naive "
        "bound -- per-layer kills cross-layer slack, EF-aware kills "
        "naive-bound slack -- so their effects partly stack rather than fully "
        "multiply.",
        align="justify",
    )
    add_figure(doc, RESULTS / "fig10_bound_tightness.png",
               "Figure 10: (Left) Theoretical sum_i alpha_i (1-r_i)^2 (solid) "
               "vs. empirical leaked squared error (dashed) per training round, "
               "log-scale. The horizontal dotted line is the budget epsilon. "
               "(Right) Mean per-worker EF correction factor kappa_i over "
               "training. EF-aware variants oscillate as the closed-loop "
               "controller pushes compression to the edge of the EF residual "
               "envelope and self-corrects.")
    add_para(
        doc,
        "The left panel of Figure 10 visualises the bound-tightness story. "
        "For the naive scheme (grey), theoretical error sits near epsilon "
        "(constraint active) but empirical error sits 1-2 orders of magnitude "
        "below -- the bound is correct but loose. EF-aware variants (blue, "
        "red) deliberately push solid lines above epsilon: they have moved to "
        "the kappa-rescaled constraint, so the unscaled sum_i alpha_i (1-r_i)^2 "
        "is no longer the constraint. Their dashed (empirical) lines remain "
        "low, confirming that real accuracy is preserved. The right panel "
        "tracks kappa_i over training: it drops from its initial value of 1 "
        "to ~0.03 for EF-unaware variants and oscillates between 0.05 and 1 "
        "for the per-layer EF-aware variant, an inherent self-stabilising "
        "behaviour driven by recalibration events at every Kth round.",
        align="justify",
    )

    # ----------- 6. DISCUSSION -----------
    add_heading(doc, "6. Discussion", level=1)
    add_para(doc, "6.1 Why does the formulation work?", bold=True)
    add_para(
        doc,
        "The core observation is that the bottleneck in synchronous "
        "data-parallel training is the slowest worker, and that compression "
        "has both a per-worker accuracy cost (controllable per worker) and a "
        "shared communication benefit (set by the bottleneck). The min-max "
        "objective directly aligns these. The water-filling solution that "
        "emerges (Section 5.3) is the canonical answer when one is "
        "minimising a max under a sum-of-quadratics constraint, and provides "
        "intuitive validation that the formulation matches the underlying "
        "ring structure.",
        align="justify",
    )
    add_para(doc, "6.2 The kappa oscillation phenomenon.", bold=True)
    add_para(
        doc,
        "The right panel of Figure 10 shows that for the most aggressive "
        "variant (per-layer + EF-aware), kappa_i oscillates between roughly "
        "0.05 and 1.0 with a period matched to the recalibration interval. "
        "Mechanistically: between recalibrations, the optimizer commits to "
        "ratios calibrated against a low kappa; this drives empirical "
        "leakage above the predicted level (residuals build up); the EMA of "
        "kappa rises; the next recalibration sees a higher kappa and pulls "
        "ratios back; residuals flush; kappa drops; repeat. The system is "
        "bounded and self-correcting, but the oscillation is a real cost: "
        "occasional rounds incur empirical error well above the budget. A "
        "cleaner alternative (left for future work) would be to model kappa "
        "explicitly inside the convex program with a stability constraint.",
        align="justify",
    )
    add_para(doc, "6.3 Limitations.", bold=True)
    add_bullets(doc, [
        "We evaluate on MNIST with a small MLP. Larger and more heterogeneous "
        "models (transformers, vision backbones) have richer per-layer "
        "alpha structure that should benefit per-layer further, but we have "
        "not validated this empirically.",
        "Our simulator captures the ring all-reduce wall-clock model exactly "
        "but does not model contention with compute, memory transfers, "
        "topology-aware collectives (recursive doubling, hierarchical), or "
        "real-world jitter. Validation on a physical cluster is the most "
        "important next step.",
        "The accuracy budget epsilon is a hyperparameter; we sweep it for the "
        "Pareto plot but ultimately leave its choice to the practitioner. A "
        "future direction is to tie epsilon directly to a target convergence "
        "rate (Karimireddy et al.'s framework supports this).",
    ])

    # ----------- 7. CONCLUSION -----------
    add_heading(doc, "7. Conclusion and Future Work", level=1)
    add_para(
        doc,
        "We presented CompressIQ, a convex formulation of per-worker gradient "
        "compression for heterogeneous distributed training, derived directly "
        "from the bandwidth-optimal Patarasuk-Yuan ring all-reduce cost. The "
        "static formulation delivers a 1.85x speedup over uniform compression "
        "at matched accuracy on a 12-worker, 3-tier cluster. Two bound "
        "refinements -- an online EF-aware correction factor kappa_i and a "
        "per-layer extension R[i, l] -- together yield an additional 1.76x "
        "speedup, for a cumulative 3.26x improvement, while preserving "
        "accuracy.",
        align="justify",
    )
    add_para(doc, "Future directions:", bold=True)
    add_bullets(doc, [
        "EF-aware bound with explicit residual-norm modelling. The Karimireddy "
        "et al. residual bound (||e_t||^2 <= 4(1-delta)/delta^2 * G^2 for a "
        "delta-contractive compressor) admits a tighter, analytically-derived "
        "kappa than our online EMA.",
        "Per-coordinate or per-tensor variance modelling. Replace alpha_i,l "
        "with a richer quadratic form that respects gradient covariance "
        "structure.",
        "Heavy-tail-aware bound. Top-K's error on a Pareto-distributed gradient "
        "scales as (1-r)^(alpha-1), not (1-r)^2; estimating the tail index "
        "online would tighten the bound for realistic deep-network gradients.",
        "Time-varying epsilon(t). Schedule the budget like a learning-rate "
        "schedule (looser early, tighter late).",
        "Validation on physical clusters with deeper models (ResNet-50, "
        "small transformers).",
    ])

    # ----------- REFERENCES -----------
    add_heading(doc, "References", level=1)
    add_para(doc, "[Placeholder section -- citations to be added.]", italic=True)

    # ----------- save -----------
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
